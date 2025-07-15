import pandas as pa
import numpy as np
from docx import Document 
document_read = pa.read_csv("2018_Central_Park_Squirrel_Census_-_Squirrel_Data_20250715.csv")
num_of_gray_color=len(document_read[document_read['Primary Fur Color'] == 'Gray'])
num_of_Black_color=len(document_read[document_read['Primary Fur Color'] == 'Black'])
num_of_Cinnamon_color=len(document_read[document_read['Primary Fur Color'] == 'Cinnamon'])
the_common_color = "gray"
num_of_pms=len(document_read[document_read['Shift']=="PM"])
num_of_ams=len(document_read[document_read['Shift']=="AM"])
the_best_time = 'PM'
num_of_adult=len(document_read[document_read['Age'] == 'Adult'])
num_of_Juvenile=len(document_read[document_read['Age'] == 'Juvenile'])
percentage_of_Juvenile = int((num_of_Juvenile/(num_of_Juvenile+num_of_adult))*100)
percentage_of_adult = int((num_of_adult/(num_of_Juvenile+num_of_adult))*100)
num_of_tree = len(document_read[document_read['Location'] == 'Above Ground'])
num_of_ground = len(document_read[document_read['Location'] == 'Ground Plane'])
the_best_location_is="The Ground"
activities = ["Running", "Chasing", "Climbing", "Eating", "Foraging"]
am_count_list=[]
am_count_list=[]
for activity in activities:
    am_count = len(document_read[(document_read['Shift'] == 'AM') & (document_read[activity] == True)])
    pm_count = len(document_read[(document_read['Shift'] == 'PM') & (document_read[activity] == True)])
doc = Document()
doc.add_heading("The Report",level=1)
doc.add_heading("Number of Gray color",level=2)
doc.add_paragraph(f"Number of Gray color squirrels: {num_of_gray_color}")
doc.add_paragraph(f"Number of Black color squirrels: {num_of_Black_color}")
doc.add_paragraph(f"Number of Cinnamon color squirrels: {num_of_Cinnamon_color}")
doc.add_paragraph(f"Most common color: {the_common_color}")

doc.add_heading("Shift Analysis", level=2)
doc.add_paragraph(f"Number of PM sightings: {num_of_pms}")
doc.add_paragraph(f"Number of AM sightings: {num_of_ams}")
doc.add_paragraph(f"Best time to spot squirrels: {the_best_time}")

doc.add_heading("Age Analysis", level=2)
doc.add_paragraph(f"Number of Adults: {num_of_adult}")
doc.add_paragraph(f"Number of Juveniles: {num_of_Juvenile}")
doc.add_paragraph(f"Percentage of Juveniles: {percentage_of_Juvenile}%")
doc.add_paragraph(f"Percentage of Adults: {percentage_of_adult}%")

doc.add_heading("Location Analysis", level=2)
doc.add_paragraph(f"Number Above Ground: {num_of_tree}")
doc.add_paragraph(f"Number on Ground Plane: {num_of_ground}")
doc.add_paragraph(f"Best location to spot squirrels: {the_best_location_is}")

doc.add_heading("Activity Analysis", level=2)
for activity in activities:
    am_count = len(document_read[(document_read['Shift'] == 'AM') & (document_read[activity] == True)])
    pm_count = len(document_read[(document_read['Shift'] == 'PM') & (document_read[activity] == True)])
    doc.add_paragraph(f"{activity}: AM count = {am_count}, PM count = {pm_count}")

# Save Word document
doc.save("Squirrel_Report.docx")

# Prepare Excel summary
summary_data = {
    "Metric": [
        "Gray Squirrels", "Black Squirrels", "Cinnamon Squirrels", "Most Common Color",
        "PM Sightings", "AM Sightings", "Best Time",
        "Adults", "Juveniles", "Juvenile %", "Adult %",
        "Above Ground", "Ground Plane", "Best Location"
    ],
    "Value": [
        num_of_gray_color, num_of_Black_color, num_of_Cinnamon_color, the_common_color,
        num_of_pms, num_of_ams, the_best_time,
        num_of_adult, num_of_Juvenile, percentage_of_Juvenile, percentage_of_adult,
        num_of_tree, num_of_ground, the_best_location_is
    ]
}

# Add activities to Excel summary
for activity in activities:
    am_count = len(document_read[(document_read['Shift'] == 'AM') & (document_read[activity] == True)])
    pm_count = len(document_read[(document_read['Shift'] == 'PM') & (document_read[activity] == True)])
    summary_data["Metric"].append(f"{activity} (AM)")
    summary_data["Value"].append(am_count)
    summary_data["Metric"].append(f"{activity} (PM)")
    summary_data["Value"].append(pm_count)

summary_df = pa.DataFrame(summary_data)
summary_df.to_excel("Squirrel_Summary.xlsx", index=False)